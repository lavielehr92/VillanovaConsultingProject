# main.py
import os
import re
import sys
import json
from pathlib import Path
from typing import Dict, Any, Tuple

from dotenv import load_dotenv
from docx import Document as DocxDocument
from openai import OpenAI

# ---------- Paths that match your project ----------
ROOT = Path(__file__).parent
ENV_PATH = ROOT / "MyKeys" / ".env"
INPUT_DIR = ROOT / "InputArticles"
OUTPUT_DIR = ROOT / "OutputArticles"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Use the exact file you showed in VS Code
DEFAULT_ARTICLE = INPUT_DIR / "These tech trends will dominate in 2023.docx"

FEATURE_KEYS = [
    "SC", "SENT",
    "INV", "INV_intens",
    "SHIP", "SHIP_intens",
    "COMM", "COMM_intens",
    "MANUF", "MANUF_intens",
    "30D", "90D"
]

SYSTEM_PROMPT = (
    "You are a careful analyst. Read the provided news article and answer EXACTLY the 12 items with the "
    "required types and ranges. Do not add extra keys."
)

USER_INSTRUCTIONS = """
Below is the text of a news article. Please read it and output a strict JSON object with these 12 fields:

- SC: boolean — Is this article broadly related to manufacturing or supply chains?
- SENT: integer from -100 (very negative) to +100 (very positive) — overall tone.
- INV: boolean — Does the article mention inventory shortages or surpluses?
- INV_intens: integer from -100 (severe shortages) to 0 (no issues) to +100 (severe gluts/surpluses).
- SHIP: boolean — Does the article mention delivery or shipping times?
- SHIP_intens: integer from -100 (very severe delays, and getting worse) to 0 (no issues) to +100 (rapidly improving speeds).
- COMM: boolean — Does the article mention commodity price changes?
- COMM_intens: integer from -100 (severe decreases) to 0 (no changes) to +100 (severe increases).
- MANUF: boolean — Does the article mention manufacturing output?
- MANUF_intens: integer from -100 (bad and decreasing) to +100 (good and increasing).
- 30D: integer from -100 (much worse) to +100 (much better) — predicted national supply-chain evolution next 30 days.
- 90D: integer from -100 (much worse) to +100 (much better) — predicted national supply-chain evolution next 90 days.

Include a field "rationale" (string, 1–4 sentences) explaining the key evidence you used.

Return ONLY a JSON object. No prose. No code fences.
"""

def read_article_text(p: Path) -> str:
    if not p.exists():
        raise FileNotFoundError(p)
    if p.suffix.lower() == ".txt":
        return p.read_text(encoding="utf-8")
    if p.suffix.lower() == ".docx":
        d = DocxDocument(str(p))
        return "\n".join(par.text for par in d.paragraphs)
    raise ValueError("Use a .txt or .docx file")

def save_results_word(out_path: Path, model: str, feat: Dict[str, Any], preview: str):
    doc = DocxDocument()
    doc.add_heading("Practice Attempt #3 — GPT Feature Extraction", 1)
    doc.add_paragraph(f"Model: {model}")

    doc.add_heading("Results (12 features)", 2)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Feature"
    table.rows[0].cells[1].text = "Value"
    order = ["SC","SENT","INV","INV_intens","SHIP","SHIP_intens",
             "COMM","COMM_intens","MANUF","MANUF_intens","30D","90D"]
    for k in order:
        r = table.add_row().cells
        r[0].text = k
        r[1].text = str(feat.get(k))

    doc.add_heading("Rationale", 2)
    doc.add_paragraph(feat.get("rationale", ""))

    doc.add_heading("Article excerpt (first 500 chars)", 2)
    doc.add_paragraph(preview[:500] + ("..." if len(preview) > 500 else ""))

    doc.save(str(out_path))

def _extract_json(text: str) -> Dict[str, Any]:
    """
    Try strict JSON first; if that fails, strip code fences or pull the largest {...} block.
    """
    t = (text or "").strip()

    # Strip ```json fences if present
    if t.startswith("```"):
        t = re.sub(r"^```(?:json)?\s*", "", t, flags=re.IGNORECASE)
        t = re.sub(r"\s*```$", "", t)

    try:
        return json.loads(t)
    except json.JSONDecodeError:
        pass

    m = re.search(r"\{.*\}", t, flags=re.DOTALL)
    if not m:
        raise RuntimeError("No JSON object found in model response.")
    return json.loads(m.group(0))

def call_gpt(model: str, article_text: str) -> Dict[str, Any]:
    """
    Uses chat.completions (works with older OpenAI SDKs). Saves raw reply and parses JSON robustly.
    """
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user",
         "content": USER_INSTRUCTIONS +
                    "\n\n--- ARTICLE START ---\n" + article_text + "\n--- ARTICLE END ---"}
    ]

    try:
        resp = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=0
        )
    except Exception as e:
        em = str(e)
        if "insufficient_quota" in em or "quota" in em.lower():
            raise RuntimeError(
                "OpenAI API error: insufficient quota. Add billing/credit or use a funded key "
                "(OPENAI_API_KEY in MyKeys/.env)."
            )
        raise

    text = (resp.choices[0].message.content or "").strip()

    # Save raw text every run for debugging
    raw_path = OUTPUT_DIR / "last_raw_response.txt"
    raw_path.write_text(text, encoding="utf-8")

    data = _extract_json(text)

    # sanity check for the 12 keys
    for k in FEATURE_KEYS:
        if k not in data:
            raise RuntimeError(f"Model JSON missing key: {k}")

    return data

def main() -> Tuple[Path, Dict[str, Any]]:
    load_dotenv(dotenv_path=ENV_PATH)

    model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
    article_path = Path(os.getenv("ARTICLE_PATH", str(DEFAULT_ARTICLE)))

    # quick sanity prints
    print("Interpreter:", sys.executable)
    print("Model:", model)
    print("API key prefix:", (os.getenv("OPENAI_API_KEY") or "")[:10] + "...")
    print("Article path:", article_path.resolve())
    print("Article exists?:", article_path.exists())

    article_text = read_article_text(article_path)

    try:
        results = call_gpt(model=model, article_text=article_text)
    except Exception as e:
        print("\nERROR:", e)
        print("If parsing failed, open this file to see the raw model reply:")
        print((OUTPUT_DIR / 'last_raw_response.txt').resolve())
        return

    out_file = OUTPUT_DIR / "practice_attempt3_results.docx"
    save_results_word(out_file, model=model, feat=results, preview=article_text)
    print(f"\nSaved: {out_file.resolve()}\n")
    print(json.dumps(results, indent=2))
    return out_file, results

if __name__ == "__main__":
    main()
